import streamlit as st
import pandas as pd
import numpy as np
import requests
import folium
from streamlit_folium import st_folium
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import tempfile 

# --- DXF KÜTÜPHANESİ KONTROLÜ ---
try:
    import ezdxf
    from ezdxf.enums import TextEntityAlignment
    HAS_EZDXF = True
except ImportError:
    HAS_EZDXF = False

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="Prefabrik Yapı Tasarım Platformu v55", layout="wide", page_icon="🏗️")
API_URL = "http://127.0.0.1:8000"

# --- TBDY METNİ ---
TBDY_TEXT = """
**8.3. MAFSALLI BAĞLANTILAR (MFB)**
* **8.3.1.4** – Yatay konumlu taşıyıcıların oturdukları mesnetlerde oluşan sürtünme kuvveti gözardı edilerek, **ideal mafsallı bağlantı (Basit Kiriş)** varsayımı yapılacaktır.
* **Tasarım Kabulü:** Kirişler, kolonlara mafsallı olarak bağlanmış olup, moment aktarmazlar.
"""

# --- FİYAT ÇEK ---
@st.cache_data
def get_default_prices():
    try:
        resp = requests.get(f"{API_URL}/varsayilan-fiyatlar", timeout=2)
        if resp.status_code == 200: return resp.json()
    except: pass
    return {'C25': 3151.28, 'C30': 3276.28, 'C35': 3401.28, 'C40': 3526.28, 'C50': 3776.28, 'celik': 21000.0, 'kalip_iscilik': 1500.0}
default_prices = get_default_prices()

# --- YARDIMCI ---
dosya_yolu = "."
afad_dosya_adi = "AAD_İller_İlceler_AFAD_Values.xlsx"
try:
    df_afad = pd.read_excel(os.path.join(dosya_yolu, afad_dosya_adi))
    iller = sorted(df_afad['il'].unique())
    zeminler = sorted(df_afad['Soil Class'].unique())
except: iller, zeminler = ["Kocaeli"], ["ZC"]

# ================== 1. TBDY DONATI HESAPLAYICI ==================
def calculate_rebar_layout(b_mm, h_mm, donati_orani):
    Ag = b_mm * h_mm
    rho_design = max(0.01, donati_orani / 100.0) 
    As_req = Ag * rho_design
    paspayi = 40 
    
    phi_list = [14, 16, 18, 20, 22, 25, 28, 32]
    valid_configs = []
    
    for phi in phi_list:
        area_one = 3.14159 * (phi/2)**2
        n = max(4, int(np.ceil(As_req / area_one)))
        if n % 2 != 0: n += 1 
        
        cevre = 2 * ((b_mm - 2*paspayi) + (h_mm - 2*paspayi))
        net_aralik = (cevre / n) - phi
        
        if net_aralik >= 40:
            symmetry_penalty = 0 if n % 4 == 0 else 1
            total_area = n * area_one
            valid_configs.append((symmetry_penalty, total_area, n, phi))
    
    if valid_configs:
        valid_configs.sort(key=lambda x: (x[0], x[1]))
        best = valid_configs[0]
        secilen_adet = best[2]
        secilen_phi = best[3]
    else:
        secilen_phi = 32
        secilen_adet = max(4, int(np.ceil(As_req / (3.14159 * 16**2))))

    corners = [(paspayi, paspayi), (b_mm-paspayi, paspayi), (b_mm-paspayi, h_mm-paspayi), (paspayi, h_mm-paspayi)]
    bars = corners[:]
    remaining = secilen_adet - 4
    
    if remaining > 0:
        len_x = b_mm - 2*paspayi; len_y = h_mm - 2*paspayi
        if remaining % 4 == 0: nx = ny = remaining // 4; add_x, add_y = True, True
        elif remaining % 2 == 0:
            if len_x >= len_y: nx = remaining // 2; ny = 0; add_x = True; add_y = False
            else: nx = 0; ny = remaining // 2; add_x = False; add_y = True
        else: nx = remaining // 2; ny = 0; add_x = True; add_y = False

        if add_x and nx > 0:
            step = len_x / (nx + 1)
            for i in range(1, nx + 1):
                bars.append((paspayi + i*step, paspayi))
                bars.append((paspayi + i*step, h_mm-paspayi))
        if add_y and ny > 0:
            step = len_y / (ny + 1)
            for i in range(1, ny + 1):
                bars.append((paspayi, paspayi + i*step))
                bars.append((b_mm-paspayi, paspayi + i*step))
                
    return secilen_adet, secilen_phi, bars

# ================== 2. DXF ÇIKTISI (V55 - ZOOM & LINETYPE FIX) ==================
def create_dxf_content(b_mm, h_mm, L_mm, bars, phi, beton_sinifi, system_info=None):
    if not HAS_EZDXF: return None
    
    # R2010 DXF Formatı
    doc = ezdxf.new('R2010') 
    msp = doc.modelspace()
    
    # --- FIX 1: DASHED LINETYPE TANIMLAMA ---
    # AutoCAD'in açılışta "DASHED not found" dememesi için
    if 'DASHED' not in doc.linetypes:
        doc.linetypes.new('DASHED', dxfattribs={
            'description': 'Dashed __ __ __',
            'pattern': [10.0, 5.0, -5.0]  # mm ölçeği için görünür pattern
        })

    # Katmanlar
    layers = [('BETON', 7), ('DONATI', 1), ('ETRIYE', 5), ('OLCU_YAZI', 2), ('SISTEM', 3), ('AKS', 8)]
    for name, color in layers:
        if name not in doc.layers: doc.layers.new(name=name, dxfattribs={'color': color})

    # ------------------ 1. KOLON PLAN (KESIT A-A) ------------------
    # Konum: (0, 0)
    msp.add_lwpolyline([(0, 0), (b_mm, 0), (b_mm, h_mm), (0, h_mm), (0, 0)], dxfattribs={'layer': 'BETON'})
    
    p = 40 
    msp.add_lwpolyline([(p, p), (b_mm-p, p), (b_mm-p, h_mm-p), (p, h_mm-p), (p, p)], dxfattribs={'layer': 'ETRIYE'})
    msp.add_line((p, p), (p+50, p+50), dxfattribs={'layer': 'ETRIYE'})
    
    for (x, y) in bars:
        msp.add_circle((x, y), phi/2, dxfattribs={'layer': 'DONATI'})
    
    # Yazı
    label_text = f"KESIT A-A: {int(b_mm)}x{int(h_mm)} | {len(bars)}%%C{phi} | C{int(beton_sinifi)}"
    msp.add_text(label_text, dxfattribs={'height': 25, 'layer': 'OLCU_YAZI'}).set_placement((0, -100))

    # ------------------ 2. KOLON BOYUNA (KESIT B-B) ------------------
    offset_x = b_mm + 1000 
    msp.add_lwpolyline([(offset_x, 0), (offset_x+h_mm, 0), (offset_x+h_mm, L_mm), (offset_x, L_mm), (offset_x, 0)], dxfattribs={'layer': 'BETON'})
    
    l_c = max(1.5 * max(b_mm, h_mm), 500, L_mm / 6)
    s_c = max(50, int(min(min(b_mm, h_mm) / 3, 150) / 10) * 10)
    s_m = max(50, int(min(min(b_mm, h_mm) / 2, 200) / 10) * 10)
    
    z = 50
    while z <= L_mm - 50:
        step = s_c if (z <= l_c or z >= L_mm - l_c) else s_m
        msp.add_line((offset_x+p, z), (offset_x+h_mm-p, z), dxfattribs={'layer': 'ETRIYE'})
        z += step
        
    msp.add_line((offset_x+p+phi/2, -600), (offset_x+p+phi/2, L_mm), dxfattribs={'layer': 'DONATI'}) 
    msp.add_line((offset_x+h_mm-p-phi/2, -600), (offset_x+h_mm-p-phi/2, L_mm), dxfattribs={'layer': 'DONATI'})
    
    msp.add_text("KESIT B-B (BOYUNA)", dxfattribs={'height': 25, 'layer': 'OLCU_YAZI'}).set_placement((offset_x, -700))
    msp.add_text(f"Etriye: %%C8/{s_c}/{s_m}", dxfattribs={'height': 25, 'layer': 'OLCU_YAZI'}).set_placement((offset_x+h_mm+50, 500))

    # Çizim sınırlarını takip et
    max_x = offset_x + h_mm + 500
    max_y = L_mm + 1000

    # ------------------ 3. SİSTEM GÖRÜNÜŞLERİ ------------------
    if system_info:
        span = system_info['aciklik'] * 1000
        k_h = system_info['kiris_h']
        aks_sayisi = system_info['aks_adedi']
        aks_aralik = system_info['aks_araligi'] * 1000
        
        # --- ÖN GÖRÜNÜŞ (FRAME) ---
        sys_x = offset_x + h_mm + 2000
        
        msp.add_lwpolyline([(sys_x, 0), (sys_x+b_mm, 0), (sys_x+b_mm, L_mm), (sys_x, L_mm), (sys_x, 0)], dxfattribs={'layer': 'SISTEM'})
        col2_x = sys_x + span - b_mm
        msp.add_lwpolyline([(col2_x, 0), (col2_x+b_mm, 0), (col2_x+b_mm, L_mm), (col2_x, L_mm), (col2_x, 0)], dxfattribs={'layer': 'SISTEM'})
        
        roof_h = span * 0.15
        mid_x = sys_x + span/2
        top_y = L_mm + k_h + roof_h
        
        beam_pts = [
            (sys_x + b_mm, L_mm), (col2_x, L_mm), (col2_x, L_mm + k_h),
            (mid_x, top_y), (sys_x + b_mm, L_mm + k_h), (sys_x + b_mm, L_mm)
        ]
        msp.add_lwpolyline(beam_pts, dxfattribs={'layer': 'SISTEM', 'color': 3})
        msp.add_text(f"SISTEM ON GORUNUS (L={int(span/1000)}m)", dxfattribs={'height': 40, 'layer': 'OLCU_YAZI'}).set_placement((sys_x, -400))

        # --- YAN GÖRÜNÜŞ (AKSLAR) ---
        side_x = sys_x + span + 4000
        msp.add_text("SISTEM YAN GORUNUS (AKSLAR)", dxfattribs={'height': 40, 'layer': 'OLCU_YAZI'}).set_placement((side_x, -400))
        
        for i in range(aks_sayisi):
            x_pos = side_x + (i * aks_aralik)
            msp.add_line((x_pos + h_mm/2, -500), (x_pos + h_mm/2, L_mm + 2500), dxfattribs={'layer': 'AKS', 'linetype': 'DASHED'})
            msp.add_text(f"AKS-{i+1}", dxfattribs={'height': 30, 'layer': 'OLCU_YAZI'}).set_placement((x_pos + h_mm/2, -600))
            msp.add_lwpolyline([(x_pos, 0), (x_pos+h_mm, 0), (x_pos+h_mm, L_mm), (x_pos, L_mm), (x_pos, 0)], dxfattribs={'layer': 'SISTEM'})
            msp.add_lwpolyline([(x_pos, L_mm), (x_pos+h_mm, L_mm), (x_pos+h_mm, L_mm+k_h), (x_pos, L_mm+k_h), (x_pos, L_mm)], dxfattribs={'layer': 'SISTEM', 'color': 2})

        total_len = (aks_sayisi-1)*aks_aralik
        msp.add_line((side_x, L_mm+k_h), (side_x + total_len + h_mm, L_mm+k_h), dxfattribs={'layer': 'SISTEM'})
        
        # Max koordinatları güncelle (Zoom için)
        max_x = side_x + total_len + 2000
        max_y = top_y + 1000

    # --- FIX 2: AUTO-ZOOM (EXTENTS & VIEWPORT) ---
    # AutoCAD açıldığında boş görünmemesi için sınırları bildiriyoruz.
    doc.header['$EXTMIN'] = (0, -2000, 0)
    doc.header['$EXTMAX'] = (max_x, max_y, 0)
    
    # Modelspace Viewport'unu ayarla (Aktif görünümü çizime odakla)
    try:
        # Merkez noktası ve yükseklik
        center_x = max_x / 2
        center_y = max_y / 2
        # VPORT tablosundaki *Active görünümü bulup güncelle
        vport = doc.viewports.get('*Active')
        if vport:
            vport.dxf.center = (center_x, center_y)
            vport.dxf.height = max_y * 1.2
    except:
        pass # Hata verirse kritik değil, EXTMAX yeterli olabilir

    # --- DOSYA KAYDETME ---
    try:
        fd, temp_path = tempfile.mkstemp(suffix='.dxf')
        os.close(fd)
        doc.saveas(temp_path)
        with open(temp_path, 'rb') as f:
            data = f.read()
        return data
    except Exception as e:
        return None
    finally:
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass

# ================== 3. 2D TEKNİK KESİT GÖRSELİ ==================
def create_2d_section_image(b_mm, h_mm, bars, phi, beton):
    fig, ax = plt.subplots(figsize=(6, 6))
    paspayi = 40
    rect = patches.Rectangle((0, 0), b_mm, h_mm, linewidth=2, edgecolor='black', facecolor='#D3D3D3')
    ax.add_patch(rect)
    etr_rect = patches.Rectangle((paspayi, paspayi), b_mm-2*paspayi, h_mm-2*paspayi, linewidth=2, edgecolor='blue', facecolor='none', linestyle='-')
    ax.add_patch(etr_rect)
    hook_len = 50 
    ax.plot([paspayi, paspayi+hook_len], [paspayi, paspayi+hook_len], color='blue', linewidth=2) 
    ax.plot([paspayi, paspayi+20], [paspayi+hook_len, paspayi+hook_len-20], color='blue', linewidth=2)
    for (x, y) in bars:
        circle = patches.Circle((x, y), radius=phi/2, edgecolor='black', facecolor='red', zorder=10)
        ax.add_patch(circle)
    if len(bars) > 4:
        ys = sorted(list(set([round(b[1],1) for b in bars])))
        xs = sorted(list(set([round(b[0],1) for b in bars])))
        if len(xs) > 2:
             for x in xs[1:-1]: ax.plot([x, x], [paspayi, h_mm-paspayi], color='blue', linestyle='--', linewidth=1)
        if len(ys) > 2:
             for y in ys[1:-1]: ax.plot([paspayi, b_mm-paspayi], [y, y], color='blue', linestyle='--', linewidth=1)
    ax.set_xlim(-50, b_mm+50); ax.set_ylim(-50, h_mm+50); ax.set_aspect('equal'); ax.axis('off')
    ax.text(b_mm/2, -30, f"{int(b_mm)} mm", ha='center', va='top', fontsize=12)
    ax.text(-30, h_mm/2, f"{int(h_mm)} mm", ha='right', va='center', rotation=90, fontsize=12)
    ax.text(b_mm/2, h_mm+20, f"{len(bars)}Φ{phi}", ha='center', va='bottom', color='red', fontsize=14, fontweight='bold')
    ax.text(b_mm/2, h_mm/2, f"C{int(beton)}", ha='center', va='center', color='gray', fontsize=20, alpha=0.3)
    buf = io.BytesIO(); plt.savefig(buf, format='png', bbox_inches='tight', dpi=150); plt.close(fig); buf.seek(0)
    return buf

# ================== 4. 3D GÖRSELLEŞTİRME ==================
def plot_3d_column_interactive(b_mm, h_mm, L_mm, beton, donati_orani):
    adet, phi, bars = calculate_rebar_layout(b_mm, h_mm, donati_orani)
    donati_metni = f"{adet}Φ{phi}"
    l_c = max(1.5 * max(b_mm, h_mm), 500, L_mm / 6)
    s_c = max(50, int(min(min(b_mm, h_mm) / 3, 150) / 10) * 10)
    s_m = max(50, int(min(min(b_mm, h_mm) / 2, 200) / 10) * 10)

    fig = go.Figure()
    x = [0, b_mm, b_mm, 0, 0, b_mm, b_mm, 0]; y = [0, 0, h_mm, h_mm, 0, 0, h_mm, h_mm]; z = [0, 0, 0, 0, L_mm, L_mm, L_mm, L_mm]
    i = [7, 0, 0, 0, 4, 4, 6, 6, 4, 0, 3, 2]; j = [3, 4, 1, 2, 5, 6, 5, 2, 0, 1, 6, 3]; k = [0, 7, 2, 3, 6, 7, 1, 1, 5, 5, 7, 6]
    fig.add_trace(go.Mesh3d(x=x, y=y, z=z, i=i, j=j, k=k, opacity=0.15, color='lightgray', name=f'Kolon C{int(beton)}', flatshading=True, hoverinfo='skip'))
    for bx, by in bars:
        fig.add_trace(go.Scatter3d(x=[bx, bx], y=[by, by], z=[0, L_mm], mode='lines', line=dict(color='red', width=phi/2.5), name=f'Boyuna {donati_metni}'))
    z_locs = []; z = 50
    while z <= l_c: z_locs.append(z); z += s_c
    z_curr = z_locs[-1] + s_m
    while z_curr <= (L_mm - l_c): z_locs.append(z_curr); z_curr += s_m
    z_curr = L_mm - l_c
    if z_curr > z_locs[-1] + s_c: z_locs.append(z_curr)
    while z_curr <= L_mm - 50: z_locs.append(z_curr); z_curr += s_c
    bulk_x, bulk_y, bulk_z = [], [], []
    paspayi = 40
    ex_x = [paspayi, b_mm-paspayi, b_mm-paspayi, paspayi, paspayi]
    ex_y = [paspayi, paspayi, h_mm-paspayi, h_mm-paspayi, paspayi]
    for z in z_locs:
        bulk_x.extend(ex_x + [None]); bulk_y.extend(ex_y + [None]); bulk_z.extend([z]*5 + [None])
    fig.add_trace(go.Scatter3d(x=bulk_x, y=bulk_y, z=bulk_z, mode='lines', line=dict(color='blue', width=2), name=f'Etriye Φ8/{int(s_c)}/{int(s_m)}'))
    max_dim = max(b_mm, h_mm); padding = max_dim * 0.8
    fig.update_layout(scene=dict(xaxis=dict(title='En', range=[-padding, b_mm+padding], showgrid=False), yaxis=dict(title='Boy', range=[-padding, h_mm+padding], showgrid=False), zaxis=dict(title='Yükseklik', range=[0, L_mm*1.1], showgrid=False), aspectmode='data'), margin=dict(l=0, r=0, b=0, t=40), title=dict(text=f"Donatı: {donati_metni}", y=0.9), legend=dict(yanchor="top", y=0.9, xanchor="left", x=0.1))
    return fig, donati_metni, f"Φ8/{int(s_c)}/{int(s_m)}", bars, phi

# ================== 5. DİĞER GÖRSELLEŞTİRME ==================
def plot_spectrum_interactive(T_kolon, SDs, SD1):
    if SDs is None: return None
    T_range = np.linspace(0, 4, 100)
    def get_sa(t): return (0.4 + 0.6*(t/(0.2*SD1/SDs)))*SDs if t < 0.2*SD1/SDs else (SDs if t <= SD1/SDs else (SD1/t if t <= 6.0 else SD1*6/(t**2)))
    Sa_vals = [get_sa(t) for t in T_range]
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=T_range, y=Sa_vals, mode='lines', name='Tasarım Spektrumu', line=dict(color='blue', width=3)))
    fig.add_trace(go.Scatter(x=[T_kolon], y=[get_sa(T_kolon)], mode='markers+text', marker=dict(size=15, color='red', symbol='x'), name='Seçilen Kolon', text=[f"T={T_kolon:.2f}s"], textposition="top right"))
    fig.update_layout(title="TBDY-2018 Tepki Spektrumu", xaxis_title="Periyot (s)", yaxis_title="Spektral İvme (g)", height=300, template="plotly_white")
    return fig

def plot_3d_system_frame(b_col, h_col, L_col, b_beam, h_beam, L_span, num_frames, spacing):
    fig = go.Figure()
    bc, hc, Lc = b_col/1000.0, h_col/1000.0, L_col/1000.0; bb, hb, Ls = b_beam/1000.0, h_beam/1000.0, L_span/1000.0; sp = spacing; roof_h = Ls * 0.15
    for n in range(num_frames):
        y = n * sp
        fig.add_trace(go.Mesh3d(x=[0,bc,bc,0,0,bc,bc,0], y=[y,y,y+bc,y+bc,y,y,y+bc,y+bc], z=[0,0,0,0,Lc,Lc,Lc,Lc], i=[7,0,0,0,4,4,6,6,4,0,3,2], j=[3,4,1,2,5,6,5,2,0,1,6,3], k=[0,7,2,3,6,7,1,1,5,5,7,6], color='gray', opacity=1, showlegend=(n==0), name="Kolon"))
        fig.add_trace(go.Mesh3d(x=[Ls-bc,Ls,Ls,Ls-bc,Ls-bc,Ls,Ls,Ls-bc], y=[y,y,y+bc,y+bc,y,y,y+bc,y+bc], z=[0,0,0,0,Lc,Lc,Lc,Lc], i=[7,0,0,0,4,4,6,6,4,0,3,2], j=[3,4,1,2,5,6,5,2,0,1,6,3], k=[0,7,2,3,6,7,1,1,5,5,7,6], color='gray', opacity=1, showlegend=False))
        # Kiriş (Üçgen Çatılı)
        kx = [0, Ls, Ls, Ls/2, 0] * 2
        ky = [y]*5 + [y+bb]*5
        # Dikdörtgen
        fig.add_trace(go.Mesh3d(x=[0, Ls, Ls, 0, 0, Ls, Ls, 0], y=[y, y, y+bb, y+bb, y, y, y+bb, y+bb], z=[Lc, Lc, Lc, Lc, Lc+hb, Lc+hb, Lc+hb, Lc+hb], i=[7, 0, 0, 0, 4, 4, 6, 6, 4, 0, 3, 2], j=[3, 4, 1, 2, 5, 6, 5, 2, 0, 1, 6, 3], k=[0, 7, 2, 3, 6, 7, 1, 1, 5, 5, 7, 6], color='#8B4513', opacity=1, showlegend=(n==0), name="Kiriş"))
        # Üçgen Çatı
        fig.add_trace(go.Mesh3d(x=[0, Ls/2, Ls/2, 0, 0, Ls/2], y=[y, y, y+bb, y+bb, y, y+bb], z=[Lc+hb, Lc+hb+roof_h, Lc+hb+roof_h, Lc+hb, Lc+hb, Lc+hb+roof_h], i=[0, 0, 0, 1], j=[1, 2, 4, 5], k=[2, 3, 5, 2], color='#A0522D', opacity=1, showlegend=False))
        fig.add_trace(go.Mesh3d(x=[Ls/2, Ls, Ls, Ls/2, Ls/2, Ls], y=[y, y, y+bb, y+bb, y, y+bb], z=[Lc+hb+roof_h, Lc+hb, Lc+hb, Lc+hb+roof_h, Lc+hb+roof_h, Lc+hb], i=[0, 0, 0, 1], j=[1, 2, 4, 5], k=[2, 3, 5, 2], color='#A0522D', opacity=1, showlegend=False))
        if n < num_frames - 1:
             for z_pos in [Lc+hb, Lc+hb+roof_h]: fig.add_trace(go.Scatter3d(x=[Ls/2, Ls/2], y=[y+bb/2, (n+1)*sp+bb/2], z=[z_pos, z_pos], mode='lines', line=dict(color='black', width=2), showlegend=False))
    fig.update_layout(scene=dict(xaxis=dict(title='Açıklık (m)'), yaxis=dict(title='Derinlik (m)'), zaxis=dict(title='Yükseklik (m)'), aspectmode='data'), margin=dict(l=0,r=0,b=0,t=30), height=400)
    return fig

# ================== RAPOR GÖRSELLERİ ==================
def create_static_3d_image(b_mm, h_mm, L_m):
    fig = plt.figure(figsize=(6, 6)); ax = fig.add_subplot(111, projection='3d')
    b, h, L = b_mm/1000.0, h_mm/1000.0, L_m
    x = [0, b, b, 0, 0, b, b, 0]; y = [0, 0, h, h, 0, 0, h, h]; z = [0, 0, 0, 0, L, L, L, L]
    faces = [[(0,0,0), (b,0,0), (b,h,0), (0,h,0)], [(0,0,L), (b,0,L), (b,h,L), (0,h,L)], [(0,0,0), (b,0,0), (b,0,L), (0,0,L)], [(b,h,0), (0,h,0), (0,h,L), (b,h,L)], [(0,0,0), (0,h,0), (0,h,L), (0,0,L)], [(b,0,0), (b,h,0), (b,h,L), (b,0,L)]]
    ax.add_collection3d(Poly3DCollection(faces, facecolors='#A9A9A9', linewidths=1.5, edgecolors='black', alpha=0.5))
    ax.set_xticks([0, b]); ax.set_yticks([0, h]); ax.set_zticks([0, L])
    ax.set_box_aspect((b, h, L)); margin = max(b, h) * 0.2; ax.set_xlim([-margin, b + margin]); ax.set_ylim([-margin, h + margin]); ax.set_zlim([0, L])
    ax.view_init(elev=25, azim=45); ax.xaxis._axinfo["grid"]['color'] =  (1,1,1,0); ax.yaxis._axinfo["grid"]['color'] =  (1,1,1,0); ax.zaxis._axinfo["grid"]['linestyle'] = ':'
    img_buf = io.BytesIO(); plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=150); plt.close(fig); img_buf.seek(0)
    return img_buf

def create_static_beam_image(b_mm, h_mm, L_mm):
    fig = plt.figure(figsize=(8, 4)); ax = fig.add_subplot(111, projection='3d')
    b, h, L = b_mm/1000.0, h_mm/1000.0, L_mm/1000.0
    x = [0, L, L, 0, 0, L, L, 0]; y = [0, 0, b, b, 0, 0, b, b]; z = [0, 0, 0, 0, h, h, h, h]
    faces = [[(0,0,0), (L,0,0), (L,b,0), (0,b,0)], [(0,0,h), (L,0,h), (L,b,h), (0,b,h)], [(0,0,0), (L,0,0), (L,0,h), (0,0,h)], [(0,b,0), (L,b,0), (L,b,h), (0,b,h)], [(0,0,0), (0,b,0), (0,b,h), (0,0,h)], [(L,0,0), (L,b,0), (L,b,h), (L,0,h)]]
    ax.add_collection3d(Poly3DCollection(faces, facecolors='brown', linewidths=1, edgecolors='black', alpha=0.6))
    ax.set_box_aspect((L, b, h)); ax.set_xlim(0, L*1.1); ax.set_ylim(0, b*2); ax.set_zlim(0, h*2)
    ax.view_init(elev=30, azim=-60); ax.xaxis._axinfo["grid"]['color'] = (1,1,1,0); ax.yaxis._axinfo["grid"]['color'] = (1,1,1,0); ax.zaxis._axinfo["grid"]['linestyle'] = ':'
    img_buf = io.BytesIO(); plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=150); plt.close(fig); img_buf.seek(0)
    return img_buf

def create_static_spectrum_image(T_kolon, SDs, SD1):
    if SDs is None: return None
    T_range = np.linspace(0, 4, 100)
    def get_sa(t): return (0.4 + 0.6*(t/(0.2*SD1/SDs)))*SDs if t < 0.2*SD1/SDs else (SDs if t <= SD1/SDs else (SD1/t if t <= 6.0 else SD1*6/(t**2)))
    Sa_vals = [get_sa(t) for t in T_range]
    fig, ax = plt.subplots(figsize=(6, 3.5)); ax.plot(T_range, Sa_vals, label='Tasarım Spektrumu', color='blue', linewidth=2); ax.scatter([T_kolon], [get_sa(T_kolon)], color='red', s=100, marker='X', label='Seçilen Kolon', zorder=5)
    ax.set_title("TBDY-2018 Tepki Spektrumu"); ax.set_xlabel("Periyot (s)"); ax.set_ylabel("Spektral İvme (g)"); ax.grid(True, linestyle='--', alpha=0.6); ax.legend()
    img_buf = io.BytesIO(); plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=150); plt.close(fig); img_buf.seek(0)
    return img_buf

def create_static_system_iso_image(b_col, h_col, L_col, b_beam, h_beam, L_span):
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.add_patch(patches.Rectangle((0, 0), b_col/1000, L_col/1000, facecolor='gray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((L_span/1000 - b_col/1000, 0), b_col/1000, L_col/1000, facecolor='gray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((0, L_col/1000), L_span/1000, h_beam/1000, facecolor='brown', edgecolor='black'))
    roof_h = L_span/1000 * 0.15
    ax.plot([0, L_span/1000/2, L_span/1000], [L_col/1000+h_beam/1000, L_col/1000+h_beam/1000+roof_h, L_col/1000+h_beam/1000], color='red', linewidth=2)
    ax.set_aspect('equal'); ax.axis('off'); ax.set_title(f"Sistem Kesiti (Açıklık: {L_span/1000}m)")
    buf = io.BytesIO(); plt.savefig(buf, format='png'); plt.close(fig); buf.seek(0)
    return buf

# --- RAPOR FONKSİYONLARI ---
def create_column_report(secilen_kolon, inputs, meta_data, fiyatlar):
    doc = Document(); head = doc.add_heading('PREFABRİK KOLON TASARIM RAPORU', 0); head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Rapor Tarihi: {pd.Timestamp.now().strftime('%d.%m.%Y %H:%M')}", style='Quote').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_heading('1. Proje Girdileri', level=1); table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
    input_data = [("İl / Zemin", f"{inputs['il']} / {inputs['zemin_sinifi']}"), ("Kolon Boyu (L)", f"{inputs['istenen_uzunluk_metre']} m"), ("Yük (P)", f"{inputs['cati_kilo']} kg"), ("Sismik (Ss-S1)", f"{meta_data['ss']:.3f} - {meta_data['s1']:.3f}")]
    for k, v in input_data: r = table.add_row().cells; r[0].text = k; r[1].text = v
    
    _, txt_boyuna, txt_etriye, bars, phi = plot_3d_column_interactive(secilen_kolon['b_mm'], secilen_kolon['h_mm'], inputs['istenen_uzunluk_metre']*1000, secilen_kolon['onrete_lass'], secilen_kolon['Donati_Orani_yuzde'])
    
    doc.add_heading('2. Seçilen Kolon', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Kesit: {int(secilen_kolon['b_mm'])}x{int(secilen_kolon['h_mm'])} mm\n").bold = True
    p.add_run(f"Beton: C{int(secilen_kolon['onrete_lass'])}\n")
    p.add_run(f"Boyuna Donatı: {txt_boyuna} (Oran: %{secilen_kolon['Donati_Orani_yuzde']:.2f})\n").font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(f"Etriye Düzeni: {txt_etriye} (TBDY Sıklaştırmalı)").font.color.rgb = RGBColor(0, 0, 255)

    doc.add_heading('2.1. Görsel Analizler', level=2)
    try:
        img_2d = create_2d_section_image(secilen_kolon['b_mm'], secilen_kolon['h_mm'], bars, phi, secilen_kolon['onrete_lass'])
        doc.add_picture(img_2d, width=Inches(3.5)); doc.add_paragraph("Şekil 1: Kolon Donatı Kesit Detayı", style='Caption')
    except: pass
    try:
        img_3d = create_static_3d_image(secilen_kolon['b_mm'], secilen_kolon['h_mm'], inputs['istenen_uzunluk_metre']); doc.add_picture(img_3d, width=Inches(3.5)); doc.add_paragraph("Şekil 2: Kolon 3D", style='Caption')
    except: pass
    
    if meta_data.get('sds'):
        try:
            img_spec = create_static_spectrum_image(secilen_kolon['Hesaplanan_Periyot_T_s'], meta_data['sds'], meta_data['sd1'])
            doc.add_picture(img_spec, width=Inches(5.5)); doc.add_paragraph(f"Şekil 3: Spektrum Eğrisi (T={secilen_kolon['Hesaplanan_Periyot_T_s']:.3f}s)", style='Caption')
        except: pass
    
    doc.add_heading('3. Analiz Sonuçları', level=1); res_table = doc.add_table(rows=1, cols=2); res_table.style = 'Table Grid'
    res_data = [("Periyot (T)", f"{secilen_kolon['Hesaplanan_Periyot_T_s']:.4f} s"), ("Moment Talebi (Md)", f"{secilen_kolon['Hesaplanan_M_Demand_kNm']:.2f} kNm"), ("Moment Kapasitesi (Mr)", f"{secilen_kolon['Hesaplanan_Mr_Kapasite_kNm']:.2f} kNm"), ("DCR Oranı", f"%{secilen_kolon['Fiziksel_Kapasite_Orani_yuzde']:.2f}"), ("Toplam Maliyet", f"{secilen_kolon['Maliyet_Endeksi']:,.2f} TL")]
    for k, v in res_data: r = res_table.add_row().cells; r[0].text = k; r[1].text = v
    doc.add_heading('4. Metraj ve Keşif', level=1)
    mt = doc.add_table(rows=1, cols=4); mt.style = 'Table Grid'
    hdr = mt.rows[0].cells; hdr[0].text='Poz'; hdr[1].text='Miktar'; hdr[2].text='Birim'; hdr[3].text='Tutar (TL)'
    vol = (secilen_kolon['Ag_mm2']/1e6) * inputs['istenen_uzunluk_metre']; w_celik = (secilen_kolon['As_mm2']/1e6) * inputs['istenen_uzunluk_metre'] * 7.85
    r1 = mt.add_row().cells; r1[0].text=f"C{int(secilen_kolon['onrete_lass'])} Beton"; r1[1].text=f"{vol:.2f} m³"; r1[2].text=f"{secilen_kolon['Kullanilan_Beton_Fiyati']:,.2f}"; r1[3].text=f"{vol*secilen_kolon['Kullanilan_Beton_Fiyati']:,.2f}"
    r2 = mt.add_row().cells; r2[0].text="İnşaat Demiri"; r2[1].text=f"{w_celik:.3f} Ton"; r2[2].text=f"{fiyatlar['fiyat_celik']:,.2f}"; r2[3].text=f"{w_celik*fiyatlar['fiyat_celik']:,.2f}"
    doc.add_paragraph(f"TOPLAM MALİYET: {secilen_kolon['Maliyet_Endeksi']:,.2f} TL").bold=True
    doc.add_heading('5. Değerlendirme', level=1); dcr = secilen_kolon['Fiziksel_Kapasite_Orani_yuzde']
    run = doc.add_paragraph("SONUÇ: GÜVENLİ." if dcr <= 100 else "SONUÇ: RİSKLİ.").runs[0]; run.font.color.rgb = RGBColor(0, 128, 0) if dcr <= 100 else RGBColor(255, 0, 0); run.bold = True; doc.add_paragraph("Kolon, fiziksel DCR kontrolünü sağlamıştır." if dcr <= 100 else "DCR %100 üzerindedir. Kesit artırımı önerilir.")
    f = io.BytesIO(); doc.save(f); f.seek(0); return f

def create_full_system_report(sistem_data, inputs, meta, fiyatlar):
    doc = Document(); doc.add_heading('SİSTEM TASARIM RAPORU', 0)
    doc.add_heading('1. Sistem Verileri', level=1); doc.add_paragraph(f"Lokasyon: {inputs['il']} ({inputs['zemin_sinifi']}) | Ss={meta['ss']:.2f}"); doc.add_paragraph(f"Toplam Yük: {inputs['total_cati_yuku_kg']} kg"); doc.add_paragraph(f"Geometri: {inputs['aks_adedi']} Aks | {inputs['aks_araligi_m']}m Aralık | {inputs['kiris_acikligi_m']}m Açıklık")
    doc.add_heading('2. Sistem Analizi', level=1)
    try:
        img = create_static_system_iso_image(sistem_data['kolon']['b_mm'], sistem_data['kolon']['h_mm'], sistem_data['kolon']['Length_mm'], sistem_data['kiris_b_mm'], sistem_data['kiris_h_mm'], inputs['kiris_acikligi_m']*1000)
        doc.add_picture(img, width=Inches(5.0)); doc.add_paragraph("Şekil 1: Sistem Kesit Görünümü", style='Caption')
    except: pass
    if meta.get('sds'):
        try:
            img_spec = create_static_spectrum_image(sistem_data['kolon']['Hesaplanan_Periyot_T_s'], meta['sds'], meta['sd1'])
            doc.add_picture(img_spec, width=Inches(5.0)); doc.add_paragraph("Şekil 2: Kolon Performansı (Spektrum)", style='Caption')
        except: pass
    doc.add_heading('3. Metraj ve Keşif Özeti', level=1)
    col = sistem_data['kolon']; k = sistem_data; n_kolon = inputs['aks_adedi'] * 2; n_kiris = inputs['aks_adedi']
    v_c = (col['Ag_mm2']/1e6) * col['Length_mm']/1000; w_c = (col['As_mm2']/1e6) * col['Length_mm']/1000 * 7.85
    v_k = (k['kiris_b_mm']/1000 * k['kiris_h_mm']/1000 * inputs['kiris_acikligi_m']); w_k = v_k * 0.02 * 7.85
    mt = doc.add_table(rows=1, cols=5); mt.style = 'Table Grid'
    hdr = mt.rows[0].cells; hdr[0].text='Eleman'; hdr[1].text='Adet'; hdr[2].text='Beton (m³)'; hdr[3].text='Demir (Ton)'; hdr[4].text='Tutar (TL)'
    r1 = mt.add_row().cells; r1[0].text='Kolonlar'; r1[1].text=str(n_kolon); r1[2].text=f"{v_c*n_kolon:.2f}"; r1[3].text=f"{w_c*n_kolon:.3f}"; r1[4].text=f"{col['Maliyet_Endeksi']*n_kolon:,.0f}"
    r2 = mt.add_row().cells; r2[0].text='Kirişler'; r2[1].text=str(n_kiris); r2[2].text=f"{v_k*n_kiris:.2f}"; r2[3].text=f"{w_k*n_kiris:.3f}"; r2[4].text=f"{k['kiris_maliyet']*n_kiris:,.0f}"
    doc.add_paragraph(f"\nGENEL TOPLAM MALİYET: {sistem_data['sistem_toplam_maliyet']:,.2f} TL").bold=True
    doc.add_heading('4. Detaylar', level=1)
    _, txt_boy_sys, txt_etr_sys, bars, phi = plot_3d_column_interactive(col['b_mm'], col['h_mm'], col['Length_mm'], col['onrete_lass'], col['Donati_Orani_yuzde'])
    doc.add_heading('4.1. Kiriş', level=2); doc.add_paragraph(f"Kesit: {int(k['kiris_b_mm'])}x{int(k['kiris_h_mm'])} mm | Beton: {k['kiris_beton']}"); doc.add_paragraph(f"Sehim: {k['kiris_sehim_mm']:.2f} mm (Limit: {k['kiris_sehim_limit_mm']:.2f} mm)")
    try: img_k = create_static_beam_image(k['kiris_b_mm'], k['kiris_h_mm'], inputs['kiris_acikligi_m']*1000); doc.add_picture(img_k, width=Inches(4.0)); doc.add_paragraph("Şekil 3: Kiriş Kesiti", style='Caption')
    except: pass
    doc.add_heading('4.2. Kolon Donatı Detayı', level=2)
    p = doc.add_paragraph(); p.add_run(f"Kesit: {int(col['b_mm'])}x{int(col['h_mm'])} mm | C{int(col['onrete_lass'])}\n"); p.add_run(f"Donatı: {txt_boy_sys}\n").font.color.rgb = RGBColor(255, 0, 0); p.add_run(f"Etriye: {txt_etr_sys}").font.color.rgb = RGBColor(0, 0, 255)
    try: img_2d = create_2d_section_image(col['b_mm'], col['h_mm'], bars, phi, col['onrete_lass']); doc.add_picture(img_2d, width=Inches(3.0)); doc.add_paragraph("Şekil 4: Kolon Donatı Kesit Detayı", style='Caption')
    except: pass
    try: img_c = create_static_3d_image(col['b_mm'], col['h_mm'], col['Length_mm']/1000); doc.add_picture(img_c, width=Inches(3.0)); doc.add_paragraph("Şekil 5: Kolon 3D", style='Caption')
    except: pass
    doc.add_heading('5. Değerlendirme', level=1)
    sehim_ok = k['kiris_sehim_mm'] <= k['kiris_sehim_limit_mm']; dcr_ok = col['Fiziksel_Kapasite_Orani_yuzde'] <= 100; p_eval = doc.add_paragraph()
    if sehim_ok and dcr_ok:
        run = p_eval.add_run("SONUÇ: GÜVENLİ."); run.font.color.rgb = RGBColor(0, 128, 0); run.bold = True
        doc.add_paragraph("Seçilen sistem hem sehim (L/200) hem de kolon kapasite (DCR < 100) şartlarını sağlamaktadır.")
    else:
        run = p_eval.add_run("SONUÇ: RİSKLİ."); run.font.color.rgb = RGBColor(255, 0, 0); run.bold = True
        if not sehim_ok: doc.add_paragraph("- Kiriş sehim limiti aşılmıştır.")
        if not dcr_ok: doc.add_paragraph("- Kolon kapasitesi yetersizdir.")
    f = io.BytesIO(); doc.save(f); f.seek(0); return f

# === ARAYÜZ ===
st.sidebar.title("🏗️ Tasarım Modu")
mod = st.sidebar.radio("Mod:", ["Kolon (Tekil)", "Sistem (Çerçeve)"])
st.sidebar.markdown("---")
tab_analiz, tab_fiyat = st.tabs(["🔍 Analiz", "💰 Fiyatlar"])
with tab_fiyat:
    st.title("Birim Fiyat Listesi")
    st.info("ℹ️ Birim fiyatlar, **T.C. Çevre, Şehircilik ve İklim Değişikliği Bakanlığı** 2025 Yılı İnşaat ve Tesisat Birim Fiyatları listesinden referans alınmıştır.")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("##### Beton Fiyatları (TL/m³)")
        p_c25 = st.number_input("C25/30", value=float(default_prices.get('C25', 3151.28))); p_c30 = st.number_input("C30/37", value=float(default_prices.get('C30', 3276.28))); p_c35 = st.number_input("C35/45", value=float(default_prices.get('C35', 3401.28))); p_c40 = st.number_input("C40/50", value=float(default_prices.get('C40', 3526.28))); p_c50 = st.number_input("C50/60", value=float(default_prices.get('C50', 3776.28)))
    with c2:
        st.markdown("##### Donatı ve İşçilik")
        p_celik = st.number_input("İnşaat Demiri (TL/ton)", value=float(default_prices.get('celik', 21000))); p_iscilik = st.number_input("İşçilik", value=float(default_prices.get('kalip_iscilik', 1500)))
    fiyat_paketi = {"fiyat_beton_c25": p_c25, "fiyat_beton_c30": p_c30, "fiyat_beton_c35": p_c35, "fiyat_beton_c40": p_c40, "fiyat_beton_c50": p_c50, "fiyat_celik": p_celik, "fiyat_iscilik": p_iscilik}
with tab_analiz:
    if mod == "Kolon (Tekil)":
        st.title("Sismik Kolon Tasarımı")
        with st.sidebar:
            st.subheader("Kolon Girdileri"); il = st.selectbox("İl", iller, index=iller.index("Kocaeli") if "Kocaeli" in iller else 0); zemin = st.selectbox("Zemin", zeminler, index=2); cati_kilo = st.number_input("Çatı Yükü (kg)", 1000); uzunluk = st.number_input("Boy (m)", min_value=1.0, value=3.5, step=0.1); hesapla_btn = st.button("Kolon Analizi Yap", type="primary")
        if hesapla_btn:
            with st.spinner("Hesaplanıyor..."):
                payload = {"il": il, "zemin_sinifi": zemin, "cati_kilo": cati_kilo, "istenen_uzunluk_metre": uzunluk, "fiyatlar": fiyat_paketi}
                try:
                    resp = requests.post(f"{API_URL}/hesapla", json=payload); data = resp.json()
                    if not data.get('uygun_kolonlar'): st.error(data.get('mesaj'))
                    else: st.session_state['kolon_sonuc'] = data; st.session_state['kolon_input'] = payload; st.session_state['secilen_kolon_index'] = 0
                except Exception as e: st.error(f"Hata: {e}")
        
        if 'kolon_sonuc' in st.session_state:
            data = st.session_state['kolon_sonuc']; df = pd.DataFrame(data['uygun_kolonlar']); payload = st.session_state['kolon_input']
            
            st.markdown("### 📊 Kolon Alternatifleri Karşılaştırması")
            table_data = []
            for i, row in df.head(10).iterrows():
                dcr_val = row['Fiziksel_Kapasite_Orani_yuzde']
                dcr_icon = "🟢" if dcr_val < 80 else ("🟡" if dcr_val < 100 else "🔴")
                table_data.append({
                    "Seçenek": f"#{i+1}",
                    "Kesit (mm)": f"{int(row['b_mm'])}x{int(row['h_mm'])}",
                    "Maliyet": f"{row['Maliyet_Endeksi']:,.0f} TL",
                    "Beton": f"C{int(row['onrete_lass'])}",
                    "Donatı (%)": f"%{row['Donati_Orani_yuzde']:.2f}",
                    "DCR (Kapasite)": f"{dcr_icon} %{dcr_val:.1f}",
                    "Periyot (T)": f"{row['Hesaplanan_Periyot_T_s']:.3f} s"
                })
            st.dataframe(pd.DataFrame(table_data), use_container_width=True, hide_index=True)
            st.divider()

            st.info("👇 Detaylı incelemek ve raporlamak için bir kolon seçiniz:")
            secenekler = []
            for i, row in df.head(10).iterrows():
                dcr = row['Fiziksel_Kapasite_Orani_yuzde']
                label = f"[{i+1}] 💰 {row['Maliyet_Endeksi']:,.0f} TL | 🏛️ {int(row['b_mm'])}x{int(row['h_mm'])} C{int(row['onrete_lass'])} (%{row['Donati_Orani_yuzde']:.1f}) | 📉 DCR: %{dcr:.1f}"
                secenekler.append(label)
            
            secilen_index = st.selectbox("Seçim:", range(len(secenekler)), format_func=lambda x: secenekler[x], index=st.session_state.get('secilen_kolon_index', 0))
            best = df.iloc[secilen_index]

            st.markdown("#### 🎯 Seçilen Kolon Özeti")
            k1, k2, k3, k4 = st.columns(4)
            with k1: st.metric("Kesit / Beton", f"{int(best['b_mm'])}x{int(best['h_mm'])}", f"C{int(best['onrete_lass'])}")
            with k2: st.metric("Maliyet", f"{best['Maliyet_Endeksi']:,.0f} TL")
            with k3: 
                dcr_val = best['Fiziksel_Kapasite_Orani_yuzde']
                st.metric("DCR Oranı", f"%{dcr_val:.1f}", delta="-Riskli" if dcr_val > 100 else "Güvenli", delta_color="inverse")
            with k4: st.metric("Donatı Oranı", f"%{best['Donati_Orani_yuzde']:.2f}", f"Periyot: {best['Hesaplanan_Periyot_T_s']:.2f}s")
            
            st.divider()

            t1, t2, t3 = st.tabs(["📝 Detaylar", "📈 Görsel Analiz", "📄 Rapor"])
            with t1:
                st.subheader("Teknik Parametreler")
                detay_veri = {
                    "Parametre": ["Boyutlar", "Beton Sınıfı", "Donatı Oranı", "Moment Kapasitesi (Mr)", "Moment Talebi (Md)", "DCR", "Periyot"],
                    "Değer": [
                        f"{int(best['b_mm'])}x{int(best['h_mm'])} mm",
                        f"C{int(best['onrete_lass'])}",
                        f"%{best['Donati_Orani_yuzde']:.2f}",
                        f"{best['Hesaplanan_Mr_Kapasite_kNm']:.2f} kNm",
                        f"{best['Hesaplanan_M_Demand_kNm']:.2f} kNm",
                        f"%{best['Fiziksel_Kapasite_Orani_yuzde']:.2f}",
                        f"{best['Hesaplanan_Periyot_T_s']:.4f} s"
                    ]
                }
                st.table(pd.DataFrame(detay_veri))

            with t2:
                col_g1, col_g2 = st.columns(2)
                with col_g1:
                    fig_spec = plot_spectrum_interactive(best['Hesaplanan_Periyot_T_s'], data['sds_degeri'], data['sd1_degeri'])
                    if fig_spec: st.plotly_chart(fig_spec, use_container_width=True)
                with col_g2:
                    fig_col, txt_boyuna, txt_etriye, bars, phi = plot_3d_column_interactive(best['b_mm'], best['h_mm'], payload['istenen_uzunluk_metre']*1000, best['onrete_lass'], best['Donati_Orani_yuzde'])
                    st.plotly_chart(fig_col, use_container_width=True); st.success(f"📌 **Seçilen Donatı:** {txt_boyuna} | **Etriye:** {txt_etriye} (Sıklaştırma/Orta)")
                    st.markdown("#### 📐 Teknik Kesit Detayı")
                    st.image(create_2d_section_image(best['b_mm'], best['h_mm'], bars, phi, best['onrete_lass']), width=400, caption="Donatı Yerleşimi (Kanca ve Çiroz Dahil)")
                    
                    dxf_data = create_dxf_content(best['b_mm'], best['h_mm'], payload['istenen_uzunluk_metre']*1000, bars, phi, best['onrete_lass'])
                    if dxf_data:
                        st.download_button(
                            label="📥 AutoCAD (.dxf) İndir",
                            data=dxf_data,
                            file_name=f"Kolon_Kesit_{int(best['b_mm'])}x{int(best['h_mm'])}.dxf",
                            mime="application/dxf",
                            type="secondary"
                        )
                    else:
                        st.warning("DXF çıktısı için 'pip install ezdxf' komutunu çalıştırın.")

                if data['secilen_konum_lat']:
                    m = folium.Map(location=[data['secilen_konum_lat'], data['secilen_konum_lon']], zoom_start=10); folium.Marker([data['secilen_konum_lat'], data['secilen_konum_lon']], popup=payload['il']).add_to(m); st_folium(m, height=200, use_container_width=True)
            with t3:
                meta = {'ss': data['ss_degeri'], 's1': data['s1_degeri'], 'sds': data.get('sds_degeri'), 'sd1': data.get('sd1_degeri', 0)}
                docx = create_column_report(best, payload, meta, fiyat_paketi)
                st.download_button("📥 Rapor İndir", docx, f"Kolon_{payload['il']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

    elif mod == "Sistem (Çerçeve)":
        st.title("Komple Sistem Tasarımı")
        with st.expander("ℹ️ Yönetmelik Referansı: TBDY-2018 8.3 (Mafsallı Bağlantılar)"): st.markdown(TBDY_TEXT)
        with st.sidebar:
            st.subheader("Sistem Girdileri"); il = st.selectbox("İl", iller); zemin = st.selectbox("Zemin", zeminler, index=2); total_yuk = st.number_input("Toplam Yük (kg)", min_value=1000, value=50000, step=1000); aks_adedi = st.number_input("Aks Sayısı", min_value=1, value=5); aciklik = st.number_input("Açıklık (m)", min_value=1.0, value=10.0); aralik = st.number_input("Aralık (m)", min_value=1.0, value=6.0); kolon_boyu = st.number_input("Kolon Boyu (m)", min_value=1.0, value=6.0);
            if st.button("Sistemi Hesapla", type="primary"):
                with st.spinner("Optimize ediliyor..."):
                    payload = {"il": il, "zemin_sinifi": zemin, "total_cati_yuku_kg": total_yuk, "kolon_boyu_m": kolon_boyu, "kiris_acikligi_m": aciklik, "aks_adedi": aks_adedi, "aks_araligi_m": aralik, "fiyatlar": fiyat_paketi}
                    try:
                        resp = requests.post(f"{API_URL}/hesapla-sistem", json=payload); data = resp.json()
                        if not data.get('uygun_sistemler'): st.error(data.get('mesaj'))
                        else: st.session_state['sis_sonuc'] = data; st.session_state['sis_input'] = payload; st.session_state['secilen_sistem_index'] = 0
                    except Exception as e: st.error(e)
        if 'sis_sonuc' in st.session_state:
            res = st.session_state['sis_sonuc']; sistemler = res['uygun_sistemler']; inp = st.session_state['sis_input']
            
            st.markdown("### 📊 Sistem Alternatifleri Karşılaştırması")
            table_data = []
            for i, s in enumerate(sistemler):
                col = s['kolon']
                dcr_val = col['Fiziksel_Kapasite_Orani_yuzde']
                dcr_icon = "🟢" if dcr_val < 80 else ("🟡" if dcr_val < 100 else "🔴")
                table_data.append({
                    "Seçenek": f"#{i+1}",
                    "Toplam Maliyet": f"{s['sistem_toplam_maliyet']:,.0f} TL",
                    "Kolon Kesit": f"{int(col['b_mm'])}x{int(col['h_mm'])}",
                    "Beton": f"C{int(col['onrete_lass'])}",
                    "Donatı (%)": f"%{col['Donati_Orani_yuzde']:.2f}",
                    "DCR (Kapasite)": f"{dcr_icon} %{dcr_val:.1f}",
                    "Kiriş Sehim": f"{s['kiris_sehim_mm']:.1f} mm"
                })
            df_compare = pd.DataFrame(table_data)
            st.dataframe(df_compare, use_container_width=True, hide_index=True)
            st.divider()

            st.info("👇 Detaylı incelemek ve raporlamak için bir sistem seçiniz:")
            secenekler_sys = []
            for i, s in enumerate(sistemler):
                c = s['kolon']
                label = f"[{i+1}] 💰 {s['sistem_toplam_maliyet']:,.0f} TL | 🏛️ {int(c['b_mm'])}x{int(c['h_mm'])} C{int(c['onrete_lass'])} (%{c['Donati_Orani_yuzde']:.1f}) | 📉 DCR: %{c['Fiziksel_Kapasite_Orani_yuzde']:.1f}"
                secenekler_sys.append(label)

            secilen_index_sys = st.selectbox("Sistem Seçimi:", range(len(secenekler_sys)), format_func=lambda x: secenekler_sys[x], index=st.session_state.get('secilen_sistem_index', 0))
            best = sistemler[secilen_index_sys]
            
            st.markdown("#### 🎯 Seçilen Sistem Özeti")
            k_met1, k_met2, k_met3, k_met4 = st.columns(4)
            with k_met1: st.metric("Toplam Maliyet", f"{best['sistem_toplam_maliyet']:,.0f} TL", delta=None)
            with k_met2:
                dcr_val = best['kolon']['Fiziksel_Kapasite_Orani_yuzde']
                st.metric("Kolon DCR", f"%{dcr_val:.2f}", delta="-Riskli" if dcr_val > 100 else "Güvenli", delta_color="inverse")
            with k_met3: st.metric("Kolon Beton/Donatı", f"C{int(best['kolon']['onrete_lass'])}", f"%{best['kolon']['Donati_Orani_yuzde']:.2f} Donatı")
            with k_met4:
                sehim_val = best['kiris_sehim_mm']; limit_val = best['kiris_sehim_limit_mm']; is_safe = sehim_val <= limit_val
                st.metric("Kiriş Sehim", f"{sehim_val:.1f} mm", f"Limit: {limit_val:.1f} mm", delta_color="normal" if is_safe else "inverse")
            st.divider()

            t1, t2 = st.tabs(["🏗️ Sistem Görünümü & Detaylar", "📄 Rapor"])
            with t1:
                st.plotly_chart(plot_3d_system_frame(best['kolon']['b_mm'], best['kolon']['h_mm'], best['kolon']['Length_mm'], best['kiris_b_mm'], best['kiris_h_mm'], inp['kiris_acikligi_m']*1000, inp['aks_adedi'], inp['aks_araligi_m']), use_container_width=True)
                c_col, c_beam = st.columns(2)
                with c_col:
                    st.subheader("🏛️ Kolon Detayları")
                    col_data = {
                        "Özellik": ["Boyutlar", "Beton Sınıfı", "Donatı Oranı", "DCR (Kapasite)", "Periyot (T)"],
                        "Değer": [
                            f"{int(best['kolon']['b_mm'])}x{int(best['kolon']['h_mm'])} mm",
                            f"C{int(best['kolon']['onrete_lass'])}",
                            f"%{best['kolon']['Donati_Orani_yuzde']:.2f}",
                            f"%{best['kolon']['Fiziksel_Kapasite_Orani_yuzde']:.2f}",
                            f"{best['kolon']['Hesaplanan_Periyot_T_s']:.3f} s"
                        ]
                    }
                    st.table(pd.DataFrame(col_data))
                    st.markdown("#### Kolon Donatı Yerleşimi (TBDY)")
                    fig_col_sys, txt_boy_sys, txt_etr_sys, bars_sys, phi_sys = plot_3d_column_interactive(best['kolon']['b_mm'], best['kolon']['h_mm'], inp['kolon_boyu_m']*1000, best['kolon']['onrete_lass'], best['kolon']['Donati_Orani_yuzde'])
                    st.plotly_chart(fig_col_sys, use_container_width=True); st.info(f"Donatı: {txt_boy_sys} | Etriye: {txt_etr_sys}")
                    st.markdown("#### 📐 Teknik Kesit")
                    st.image(create_2d_section_image(best['kolon']['b_mm'], best['kolon']['h_mm'], bars_sys, phi_sys, best['kolon']['onrete_lass']), width=400, caption="Donatı Yerleşimi")
                    
                    sys_info_dxf = {
                        'aciklik': inp['kiris_acikligi_m'],
                        'kiris_h': best['kiris_h_mm'],
                        'kiris_b': best['kiris_b_mm'],
                        'aks_adedi': inp['aks_adedi'],
                        'aks_araligi': inp['aks_araligi_m']
                    }
                    dxf_data_sys = create_dxf_content(best['kolon']['b_mm'], best['kolon']['h_mm'], inp['kolon_boyu_m']*1000, bars_sys, phi_sys, best['kolon']['onrete_lass'], sys_info_dxf)
                    if dxf_data_sys:
                        st.download_button(
                            label="📥 Tam Sistem AutoCAD (.dxf) İndir",
                            data=dxf_data_sys,
                            file_name=f"Sistem_Projesi_{inp['il']}.dxf",
                            mime="application/dxf",
                            type="secondary"
                        )

                with c_beam:
                    st.subheader("➖ Kiriş Detayları")
                    beam_data = {
                        "Özellik": ["Boyutlar", "Beton Sınıfı", "Sehim (Gerçekleşen)", "Sehim Limiti (L/200)", "Maliyet"],
                        "Değer": [
                            f"{int(best['kiris_b_mm'])}x{int(best['kiris_h_mm'])} mm",
                            f"{best['kiris_beton']}",
                            f"{best['kiris_sehim_mm']:.2f} mm",
                            f"{best['kiris_sehim_limit_mm']:.2f} mm",
                            f"{best['kiris_maliyet']:,.0f} TL"
                        ]
                    }
                    st.table(pd.DataFrame(beam_data))
                c_spec, c_map = st.columns(2)
                with c_spec:
                    st.markdown("#### Kolon Spektrum Performansı")
                    fig_spec = plot_spectrum_interactive(best['kolon']['Hesaplanan_Periyot_T_s'], res['sds_degeri'], res['sd1_degeri'])
                    if fig_spec: st.plotly_chart(fig_spec, use_container_width=True)
                with c_map:
                    if res['secilen_konum_lat']:
                        st.markdown("#### Proje Konumu")
                        m = folium.Map(location=[res['secilen_konum_lat'], res['secilen_konum_lon']], zoom_start=10); folium.Marker([res['secilen_konum_lat'], res['secilen_konum_lon']], popup=inp['il']).add_to(m); st_folium(m, height=250, use_container_width=True)
            with t2:
                meta = {'ss': res['ss_degeri'], 's1': res['s1_degeri'], 'sds': res.get('sds_degeri'), 'sd1': res.get('sd1_degeri')}
                docx = create_full_system_report(best, inp, meta, fiyat_paketi)
                st.download_button("📥 Sistem Raporu İndir", docx, "Sistem.docx")